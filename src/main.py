from __future__ import annotations

import argparse
import json
from dataclasses import asdict
from pathlib import Path

from generator import generate_draft
from journal_adapter import get_journal_rules, list_supported_journals
from loader import load_reference_documents
from preprocess import clean_text
from prompt_builder import build_drafting_context, build_query, build_structured_prompt
from schemas import DraftOutput, UserInput


def parse_list(value: str) -> list[str]:
    return [item.strip() for item in value.split(";") if item.strip()]


def load_user_input_from_config(config_path: Path) -> UserInput:
    config_path = config_path.expanduser().resolve()
    data = json.loads(config_path.read_text(encoding="utf-8"))
    reference_folder = Path(data.get("reference_docs_folder", "data/references")).expanduser()
    if not reference_folder.is_absolute():
        reference_folder = config_path.parent / reference_folder

    return UserInput(
        journal=data["journal"],
        topic=data["topic"],
        main_findings=list(data.get("main_findings", [])),
        methods=list(data.get("methods", [])),
        keywords=list(data.get("keywords", [])),
        reference_docs_folder=str(reference_folder),
        writing_instructions=str(data.get("writing_instructions", "")),
        top_k=int(data.get("top_k", 3)),
    )


def load_config(config_path: Path) -> dict:
    return json.loads(config_path.expanduser().resolve().read_text(encoding="utf-8"))


def save_json(output: DraftOutput, path: Path) -> None:
    path.write_text(json.dumps(output.to_dict(), indent=2, ensure_ascii=False), encoding="utf-8")


def save_markdown(output: DraftOutput, path: Path) -> None:
    outline = "\n".join(f"- {item}" for item in output.outline)
    retrieved = "\n".join(f"- {name}" for name in output.retrieved_documents) or "- None"
    markdown = f"""# Scientific Draft

## Title
{output.title}

## Abstract
{output.abstract}

## Paper Outline
{outline}

## Retrieved Documents
{retrieved}

## Structured Prompt for a Future Local LLM
```text
{output.prompt}
```
"""
    path.write_text(markdown, encoding="utf-8")


def save_pdf(output: DraftOutput, path: Path) -> None:
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError as exc:
        raise ImportError(
            "PDF export requires reportlab. Install project dependencies with: "
            "pip install -r requirements.txt"
        ) from exc

    styles = getSampleStyleSheet()
    document = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        title="Scientific Draft",
        rightMargin=54,
        leftMargin=54,
        topMargin=54,
        bottomMargin=54,
    )
    story = [
        Paragraph("Scientific Draft", styles["Title"]),
        Spacer(1, 12),
        Paragraph("Title", styles["Heading2"]),
        Paragraph(output.title, styles["BodyText"]),
        Spacer(1, 12),
        Paragraph("Abstract", styles["Heading2"]),
        Paragraph(output.abstract, styles["BodyText"]),
        Spacer(1, 12),
        Paragraph("Paper Outline", styles["Heading2"]),
    ]

    for item in output.outline:
        story.append(Paragraph(item, styles["BodyText"]))
        story.append(Spacer(1, 6))

    story.extend(
        [
            Spacer(1, 12),
            Paragraph("Retrieved Documents", styles["Heading2"]),
        ]
    )
    if output.retrieved_documents:
        for document_name in output.retrieved_documents:
            story.append(Paragraph(document_name, styles["BodyText"]))
            story.append(Spacer(1, 4))
    else:
        story.append(Paragraph("None", styles["BodyText"]))

    document.build(story)


def save_docx(output: DraftOutput, path: Path) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(
            "Word export requires python-docx. Install project dependencies with: "
            "pip install -r requirements.txt"
        ) from exc

    document = Document()
    document.add_heading("Scientific Draft", level=1)

    document.add_heading("Title", level=2)
    document.add_paragraph(output.title)

    document.add_heading("Abstract", level=2)
    document.add_paragraph(output.abstract)

    document.add_heading("Paper Outline", level=2)
    for item in output.outline:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Retrieved Documents", level=2)
    if output.retrieved_documents:
        for document_name in output.retrieved_documents:
            document.add_paragraph(document_name, style="List Bullet")
    else:
        document.add_paragraph("None")

    document.add_page_break()
    document.add_heading("Structured Prompt for a Future Local LLM", level=2)
    document.add_paragraph(output.prompt)

    document.save(path)


def inspect_references(reference_folder: str | Path) -> dict:
    docs = load_reference_documents(reference_folder)
    summary = {
        "reference_folder": str(Path(reference_folder).expanduser().resolve()),
        "document_count": len(docs),
        "documents": [],
    }
    for doc in docs:
        cleaned = clean_text(doc.content)
        summary["documents"].append(
            {
                "filename": doc.filename,
                "characters": len(doc.content),
                "cleaned_characters": len(cleaned),
                "has_text": bool(cleaned.strip()),
            }
        )
    return summary


def run_pipeline(
    user_input: UserInput,
    output_dir: Path,
    model_name: str,
    generation_backend: str = "template",
    llama_model: str = "llama3.1",
    ollama_url: str = "http://127.0.0.1:11434",
    fallback_to_template: bool = True,
) -> DraftOutput:
    from embedder import embed_documents, embed_query, load_embedding_model
    from retriever import retrieve_top_k

    docs = load_reference_documents(user_input.reference_docs_folder)
    for doc in docs:
        doc.cleaned_content = clean_text(doc.content)

    model = load_embedding_model(model_name)
    embed_documents(docs, model)

    query = build_query(user_input)
    query_embedding = embed_query(query, model)
    retrieved_docs = retrieve_top_k(docs, query_embedding, user_input.top_k)

    journal_rules = get_journal_rules(user_input.journal)
    context = build_drafting_context(user_input, journal_rules, retrieved_docs)
    prompt = build_structured_prompt(context)

    if generation_backend == "template":
        output = generate_draft(context, prompt)
    elif generation_backend == "ollama":
        try:
            from llm_generator import generate_draft_with_ollama

            output = generate_draft_with_ollama(
                context=context,
                prompt=prompt,
                model=llama_model,
                ollama_url=ollama_url,
            )
        except Exception as exc:
            if not fallback_to_template:
                raise
            print(f"Warning: local LLM generation failed, using template fallback. {exc}")
            output = generate_draft(context, prompt)
    else:
        raise ValueError(f"Unsupported generation backend: {generation_backend}")

    output_dir.mkdir(parents=True, exist_ok=True)
    save_json(output, output_dir / "draft_output.json")
    save_markdown(output, output_dir / "draft_output.md")
    save_pdf(output, output_dir / "draft_output.pdf")
    save_docx(output, output_dir / "draft_output.docx")

    return output


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Local retrieval-based scientific title, abstract, and outline assistant."
    )
    parser.add_argument("--config", type=Path, default=Path("config.json"))
    parser.add_argument("--output-dir", type=Path, default=Path("outputs"))
    parser.add_argument("--model", default="sentence-transformers/all-MiniLM-L6-v2")
    parser.add_argument("--generation-backend", choices=["template", "ollama"])
    parser.add_argument("--llama-model")
    parser.add_argument("--ollama-url")
    parser.add_argument("--no-template-fallback", action="store_true")
    parser.add_argument("--inspect-references", action="store_true")
    parser.add_argument("--list-journals", action="store_true")
    parser.add_argument("--journal")
    parser.add_argument("--topic")
    parser.add_argument("--main-findings", help="Semicolon-separated list.")
    parser.add_argument("--methods", help="Semicolon-separated list.")
    parser.add_argument("--keywords", help="Semicolon-separated list.")
    parser.add_argument("--writing-instructions")
    parser.add_argument("--references", help="Folder containing .md and .txt reference documents.")
    parser.add_argument("--top-k", type=int)
    return parser


def main() -> None:
    parser = build_arg_parser()
    args = parser.parse_args()

    if args.list_journals:
        print("\n".join(list_supported_journals()))
        return

    config_data = {}
    if args.config.exists():
        config_data = load_config(args.config)

    if args.inspect_references:
        reference_folder = args.references or config_data.get("reference_docs_folder", "data/references")
        if not Path(reference_folder).expanduser().is_absolute():
            reference_folder = args.config.expanduser().resolve().parent / reference_folder
        print(json.dumps(inspect_references(reference_folder), indent=2, ensure_ascii=False))
        return

    if args.topic:
        user_input = UserInput(
            journal=args.journal or "Generic Scientific Journal",
            topic=args.topic,
            main_findings=parse_list(args.main_findings or ""),
            methods=parse_list(args.methods or ""),
            keywords=parse_list(args.keywords or ""),
            reference_docs_folder=args.references or "data/references",
            writing_instructions=args.writing_instructions or "",
            top_k=args.top_k or 3,
        )
    else:
        user_input = load_user_input_from_config(args.config)

    generation_backend = args.generation_backend or config_data.get("generation_backend", "template")
    llama_model = args.llama_model or config_data.get("llama_model", "llama3.1")
    ollama_url = args.ollama_url or config_data.get("ollama_url", "http://127.0.0.1:11434")

    output = run_pipeline(
        user_input=user_input,
        output_dir=args.output_dir,
        model_name=args.model,
        generation_backend=generation_backend,
        llama_model=llama_model,
        ollama_url=ollama_url,
        fallback_to_template=not args.no_template_fallback,
    )
    print(json.dumps({"saved": str(args.output_dir), **asdict(output)}, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
